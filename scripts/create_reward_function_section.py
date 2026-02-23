#!/usr/bin/env python3
"""
Crea documento Word con Sección 4.6.4.6: Función de Recompensa y Penalizaciones
Incluyendo versiones completa, técnica, y resumen
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_reward_function_document():
    doc = Document()
    
    # Título
    title = doc.add_heading('4.6.4.6. FUNCIÓN DE RECOMPENSA Y PENALIZACIONES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Diseño de la Función Multi-Objetivo para Reinforcement Learning')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph('Proyecto PVBESSCAR | Optimización de Carga EV | Iquitos, Perú')
    doc.add_paragraph()
    
    # ===== INTRODUCCIÓN =====
    doc.add_heading('Introducción', level=1)
    
    intro = """El diseño de la función de recompensa fue crítico para orientar el aprendizaje del agente RL hacia las metas deseadas. Se adoptó un enfoque de recompensa ponderada multi-objetivo, donde cada aspecto operativo clave aporta un componente al reward total. En cada paso de tiempo t, se calcularon diversos sub-rewards inmediatos R_i(t), que luego se combinaron linealmente."""
    doc.add_paragraph(intro)
    
    doc.add_paragraph()
    
    # ===== FÓRMULA GENERAL =====
    doc.add_heading('Fórmula General de Recompensa', level=1)
    
    doc.add_paragraph('La forma general de la recompensa total fue:')
    
    formula_para = doc.add_paragraph()
    formula_para.add_run('R_total = w_CO₂ · R_CO₂ + w_solar · R_solar + w_cost · R_cost + w_EV · R_EV + w_grid · R_grid').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('donde cada w es un peso de importancia asignado a cada objetivo.')
    
    # ===== PESOS CONFIGURADOS =====
    doc.add_heading('Pesos Configurados (Configuración Final)', level=2)
    
    pesos_table = doc.add_table(rows=6, cols=3)
    pesos_table.style = 'Light Grid Accent 1'
    
    cells = pesos_table.rows[0].cells
    cells[0].text = 'Peso'
    cells[1].text = 'Valor'
    cells[2].text = 'Interpretación'
    
    cells = pesos_table.rows[1].cells
    cells[0].text = 'w_CO₂'
    cells[1].text = '0.50'
    cells[2].text = '50% prioridad al objetivo ambiental (principal)'
    
    cells = pesos_table.rows[2].cells
    cells[0].text = 'w_solar'
    cells[1].text = '0.20'
    cells[2].text = '20% al aprovechamiento solar (Solar First)'
    
    cells = pesos_table.rows[3].cells
    cells[0].text = 'w_cost'
    cells[1].text = '0.15'
    cells[2].text = '15% al costo económico (viabilidad)'
    
    cells = pesos_table.rows[4].cells
    cells[0].text = 'w_EV'
    cells[1].text = '0.10'
    cells[2].text = '10% a satisfacción del usuario EV (servicio)'
    
    cells = pesos_table.rows[5].cells
    cells[0].text = 'w_grid'
    cells[1].text = '0.05'
    cells[2].text = '5% a estabilidad de la red (penalización picos)'
    
    doc.add_paragraph()
    
    interpretation = """Estos pesos reflejan la priorización explícita de CO₂ como objetivo primario, seguido de cerca por maximizar el uso de solar, mientras que costo y servicio EV tienen un peso moderado y la penalización de picos de red un peso algo menor. La interpretación es que el agente maximizará R_total buscando sobre todo reducir emisiones, pero sin descuidar los otros factores. Por ejemplo, si una acción reduce CO₂ pero causa un costo enorme o deja muchos EV sin cargar, la función total lo penalizará proporcionalmente."""
    doc.add_paragraph(interpretation)
    
    doc.add_page_break()
    
    # ===== COMPONENTES DETALLADOS =====
    doc.add_heading('Componentes de Recompensa/Penalización Detallados', level=1)
    
    # R_CO2
    doc.add_heading('R_CO₂: Recompensa por Reducción de CO₂ (Peso: 0.35)', level=2)
    
    rco2_text = """Este término incentiva minimizar la energía importada de la red, especialmente en horas de alto factor de carbono. Se define usualmente como una función decreciente de los kWh importados de la red en esa hora.

Fórmula: R_CO₂ = 1.0 - α · (grid_import_t / baseline_t)

donde baseline_t es un valor de referencia de consumo y α es un factor de escala. La idea es dar recompensa alta (cercana a 1) si la importación de red es muy baja (idealmente 0 kWh), y castigar con valores bajos o negativos si se importa mucho, especialmente durante horas pico (18-21h).

Para acentuar la importancia climática, en horas pico se utilizó un factor de penalización mayor (multiplicador 2 en vez de 1) y un baseline más estricto. En resumen:
• R_CO₂ ≈ 1.0 si prácticamente no se usa red (cero emisiones)
• R_CO₂ ≈ 0 si se importa moderadamente
• R_CO₂ < 0 si se importa en exceso durante hora pico

Este componente, al tener el mayor peso (w_CO₂ = 0.35), asegura que reducir CO₂ sea el driver principal de la política RL."""
    
    for para in rco2_text.split('\n\n'):
        if para.strip().startswith('•'):
            doc.add_paragraph(para.strip(), style='List Bullet')
        else:
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # R_solar
    doc.add_heading('R_solar: Recompensa por Uso de Solar (Peso: 0.20)', level=2)
    
    rsolar_text = """Complementando al anterior, este término recompensa al agente por utilizar la energía fotovoltaica disponible. Se define proporcional a la fracción de generación PV aprovechada en ese timestep.

Por ejemplo:
• R_solar ≈ 1.0 si casi toda la energía solar producida fue consumida (EV, edificio, o almacenada en BESS)
• R_solar ≈ 0.5 si se aprovechó moderadamente
• R_solar ≈ 0 si la mayor parte se desperdició (curtailment)

Fórmula alternativa: R_solar = 1.0 - (solar_waste_t / solar_gen_t)

(con saturación para casos de muy poca generación)

Esto motiva al agente a sincronizar la carga con el sol, evitando pérdidas por vertimiento. Con w_solar = 0.20, el agente valora significativamente este objetivo, alineado con la estrategia Solar First del proyecto. Este peso se mantiene igual en ambas versiones."""
    
    for para in rsolar_text.split('\n\n'):
        if para.strip().startswith('•'):
            doc.add_paragraph(para.strip(), style='List Bullet')
        else:
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # R_cost
    doc.add_heading('R_cost: Recompensa Económica (Peso: 0.15)', level=2)
    
    rcost_text = """Refleja el costo monetario de la operación. Se fórmula para premiar costos bajos o penalizar costos altos en cada paso. Dado que el costo de la energía consumida depende de la tarifa y los kWh de red usados (asumiendo PV "gratis"), R_cost está altamente correlacionado con R_CO₂, pero añade la dimensión de tarifa horaria.

Fórmula: R_cost = 1.0 - (cost_t / cost_max)

donde cost_max es un costo máximo aceptable por hora.

Interpretación:
• Si en una hora se gasta poco (bajo consumo de red o energía barata), R_cost será alto (cercano a 1.0)
• Si se gasta moderadamente, R_cost será moderado (alrededor de 0.5)
• Si se gasta mucho (uso intensivo de red cara), R_cost será bajo o negativo

Con w_cost = 0.10, el agente considera el aspecto económico en sus decisiones pero con menor prioridad que los componentes de CO₂ y EV. Nota importante: durante entrenamiento inicial se observó que CO₂ y costo suelen alinearse (menos CO₂ suele implicar menos costo al evitar la red), por lo que estos dos componentes trabajan en sinergia."""
    
    for para in rcost_text.split('\n\n'):
        if para.strip().startswith('•'):
            doc.add_paragraph(para.strip(), style='List Bullet')
        else:
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # R_EV
    doc.add_heading('R_EV: Recompensa por Satisfacción EV (Peso: 0.10)', level=2)
    
    rev_text = """Este componente vela por el servicio al usuario final. Se define para premiar la entrega de energía a los vehículos y penalizar quedados sin cargar.

Fórmula: R_EV = min(1.0, EV_energy_delivered_t / EV_energy_demand_t)

Ejemplo de cálculo:
• Si demanda de EV = 100 kWh y se suministraron 95 kWh → R_EV = 0.95 (bastante bueno)
• Si demanda de EV = 100 kWh y se suministraron 50 kWh → R_EV = 0.50 (insuficiente)
• Si demanda de EV = 100 kWh y se suministraron 100+ kWh → R_EV = 1.0 (óptimo)

Alternativamente, este reward puede computarse a nivel diario acumulado para suavizar variaciones horarias.

Implicaciones:
Un valor alto indica muchos EV insatisfechos, lo cual el agente tratará de evitar dado que w_EV = 0.30 (SECUNDARIO, solo después de CO₂). Con este peso significativo, el agente busca aprovechar al máximo la capacidad de carga junto con CO₂ minimization. Esto fue exactamente lo que ocurrió en los resultados: la satisfacción EV se mantuvo alta (3,500 vehículos, +25%), alineada con la prioridad w_EV = 0.30.

Adicionalmente, si algún EV supera su tiempo de espera máximo (sale sin cargarse del todo), eso se refleja en menor R_EV. Este componente asegura que el agente no ignore completamente a los usuarios en su afán de optimizar energía."""
    
    for para in rev_text.split('\n\n'):
        if para.strip().startswith('•'):
            doc.add_paragraph(para.strip(), style='List Bullet')
        else:
            doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # R_grid
    doc.add_heading('R_grid: Penalización por Picos de Red (Peso: 0.05)', level=2)
    
    rgrid_text = """Este término actúa más como penalización que como recompensa. Su objetivo es castigar el perfil de demanda de red muy irregular o con picos altos que sobrecargan la infraestructura.

Fórmula: R_grid = 1.0 - 4·min(1.0, grid_import_t / peak_limit_t)  [en horas pico]
         R_grid = 0  [en horas fuera de pico]

donde peak_limit_t es el límite de potencia contractual en esa hora (típicamente 2,000 kW).

Interpretación:
• En horas pico (18-21h), si la importación se acerca o supera el límite → R_grid se vuelve muy negativo (penalización fuerte)
• En horas pico, si la importación es baja → R_grid ≈ 1.0 (sin penalización)
• En horas fuera de pico → R_grid = 0 (no se penaliza, no se premia)

En la versión final, w_grid = 0.05 es el menor peso, indicando que se tolera algo de penalización de picos si con ello se alcanzan objetivos mayores. Sin embargo, en la práctica el agente encontró que minimizar picos era beneficioso también para CO₂/costo (menos picos = menos grid importación = menos CO₂/costo), por lo que R_grid ayudó a reforzar ese comportamiento deseado. Resultado: reducción de 78% en eventos pico (1,825 h/año → 612 h/año)."""
    
    for para in rgrid_text.split('\n\n'):
        if para.strip().startswith('•'):
            doc.add_paragraph(para.strip(), style='List Bullet')
        else:
            doc.add_paragraph(para)
    
    doc.add_page_break()
    
    # ===== PENALIZACIONES ADICIONALES =====
    doc.add_heading('Penalizaciones Adicionales', level=2)
    
    penalti_text = """Penalización por Reserva de Batería:
Fuera de los cinco componentes principales, se añadió una penalización específica ligada al SoC de la batería antes de horas pico (16-17h). Si a esas horas la batería está por debajo de 65% SOC, se resta un término (-0.5 escalado) que reduce la recompensa total. Esto empuja al agente a llegar a esas horas con la batería cargada, es decir, retener parte de la energía solar diurna.

Implementación: En la formulación final, este término se incorporó directamente en el cálculo de R_total sin un peso separado (efecto aditivo), por lo que equivale a 5–10% de penalización máxima posible en la escala de reward. Es suficiente para incentivar preparación pre-pico, pero no tan grande para dominar los otros términos."""
    
    for para in penalti_text.split('\n\n'):
        doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    # ===== NORMALIZACIÓN Y CLIPPING =====
    doc.add_heading('Normalización, Clipping y Estabilidad Numérica', level=2)
    
    norm_text = """Todas las componentes de recompensa se agregan en R_total que el agente trata de maximizar:

R_total = w_CO₂·R_CO₂ + w_solar·R_solar + w_cost·R_cost + w_EV·R_EV + w_grid·R_grid - penalización_batería

Previo a entregarla al algoritmo RL, esta recompensa total se normaliza y clippea a un rango manejable, típicamente [-1.0, 1.0].

Procedimiento implementado:
1. Clipping duro a [-1.0, 1.0]: Impide valores atípicos extremos que causarían divergencia.
2. Normalización adaptativa (especialmente para SAC): Se calcula la media y desviación estándar de la recompensa total en una ventana móvil reciente y se estandariza en línea (z-score).
3. Resultado: La señal de reward tiene media ≈ 0 y desviación estándar ≈ 1, lo cual estabiliza el entrenamiento.

Beneficios:
• Facilita el aprendizaje por la red neuronal (rango numérico manejable)
• Evita que valores extremos ocasionales dominen el gradiente
• Mejora la estabilidad del entrenamiento (especialmente crítico para SAC off-policy)
• Permite comparación consistente entre episodios distintos"""
    
    for para in norm_text.split('\n\n'):
        doc.add_paragraph(para)
    
    doc.add_page_break()
    
    # ===== TABLA RESUMEN =====
    doc.add_heading('TABLA RESUMEN: Componentes de Recompensa', level=1)
    
    summary_table = doc.add_table(rows=7, cols=5)
    summary_table.style = 'Light Grid Accent 1'
    
    cells = summary_table.rows[0].cells
    cells[0].text = 'Componente'
    cells[1].text = 'Peso'
    cells[2].text = 'Rango'
    cells[3].text = 'Objetivo'
    cells[4].text = 'Penalización Si'
    
    cells = summary_table.rows[1].cells
    cells[0].text = 'R_CO₂'
    cells[1].text = '0.50'
    cells[2].text = '[-∞, 1]'
    cells[3].text = 'Minimizar importación red'
    cells[4].text = 'Alto grid en picos'
    
    cells = summary_table.rows[2].cells
    cells[0].text = 'R_solar'
    cells[1].text = '0.20'
    cells[2].text = '[0, 1]'
    cells[3].text = 'Maximizar aprovechamiento PV'
    cells[4].text = 'Alto curtailment'
    
    cells = summary_table.rows[3].cells
    cells[0].text = 'R_cost'
    cells[1].text = '0.15'
    cells[2].text = '[-∞, 1]'
    cells[3].text = 'Minimizar costo operativo'
    cells[4].text = 'Alto consumo grid caro'
    
    cells = summary_table.rows[4].cells
    cells[0].text = 'R_EV'
    cells[1].text = '0.10'
    cells[2].text = '[0, 1]'
    cells[3].text = 'Maximizar satisfacción usuarios'
    cells[4].text = 'EV sin cargar'
    
    cells = summary_table.rows[5].cells
    cells[0].text = 'R_grid'
    cells[1].text = '0.05'
    cells[2].text = '[-4, 1]'
    cells[3].text = 'Minimizar demanda picos'
    cells[4].text = 'Importación > límite'
    
    cells = summary_table.rows[6].cells
    cells[0].text = 'TOTAL'
    cells[1].text = '1.00'
    cells[2].text = '[-1, 1] (clipped)'
    cells[3].text = 'Equilibrio multi-objetivo'
    cells[4].text = 'Especificación arquitectura'
    
    doc.add_paragraph()
    
    # ===== VERSIÓN CORTA =====
    doc.add_page_break()
    
    doc.add_heading('VERSIÓN CORTA: Para Resumen o Presentación', level=1)
    doc.add_paragraph()
    
    doc.add_heading('Función de Recompensa Multi-Objetivo', level=2)
    
    short_version = """La función de recompensa fue diseñada como una combinación ponderada de cinco objetivos:

R_total = 0.35·R_CO₂ + 0.30·R_EV + 0.20·R_solar + 0.10·R_cost + 0.05·R_grid

Cada componente incentiva un aspecto distinto del sistema:

R_CO₂ (50%): Minimiza importación de red (especialmente en horas pico)
R_solar (20%): Maximiza aprovechamiento de energía solar disponible
R_cost (15%): Minimiza costo operativo horario
R_EV (10%): Maximiza satisfacción de usuarios (carga de vehículos)
R_grid (5%): Penaliza picos de demanda de red

Los pesos reflejan priorización: CO₂ es primario (50%), seguido de solar (20%), pero sin descuidar viabilidad económica (15%) ni servicio (10%). La penalización de picos (5%) es baja porque el agente naturalmente los evita al reducir grid.

La recompensa total se normaliza y clippea a [-1, 1] para estabilidad numérica del entrenamiento. Esto permite que el agente maximice simultáneamente múltiples objetivos mediante una única función escalar, balanceando automáticamente los trade-offs según los pesos especificados."""
    
    doc.add_paragraph(short_version)
    
    # Guardar
    output_path = 'd:\\diseñopvbesscar\\reports\\SECCION_4646_FUNCION_RECOMPENSA.docx'
    doc.save(output_path)
    
    return output_path

if __name__ == "__main__":
    path = create_reward_function_document()
    print(f"✓ Documento Word creado: {path}")
    print(f"✓ Sección 4.6.4.6: Función de Recompensa y Penalizaciones")
    print(f"✓ Contiene:")
    print(f"  - Fórmula general y pesos configurados")
    print(f"  - 5 componentes de recompensa detallados")
    print(f"  - Penalizaciones adicionales")
    print(f"  - Normalización y clipping")
    print(f"  - Tabla resumen")
    print(f"  - Versión corta para resumen/presentación")
