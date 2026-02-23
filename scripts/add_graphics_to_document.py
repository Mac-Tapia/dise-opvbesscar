#!/usr/bin/env python3
"""
Incorporar gráficos PNG en el documento integrado (5.2-5.5)
Inserta imágenes de alta resolución en las secciones correspondientes
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).parent.parent))

def add_graphics_to_document():
    """Añadir gráficos al documento integrado"""
    
    print("\n" + "="*80)
    print("INCORPORANDO GRÁFICOS EN DOCUMENTO INTEGRADO")
    print("="*80 + "\n")
    
    # Cargar documento integrado existente
    doc_path = Path('TESIS_SECCIONES_5_2_a_5_5_INTEGRADO_COMPLETO.docx')
    if not doc_path.exists():
        # Intentar en outputs
        doc_path = Path('outputs/TESIS_SECCIONES_5_2_a_5_5_INTEGRADO_COMPLETO.docx')
    
    if not doc_path.exists():
        print(f"✗ Documento no encontrado en ninguna ubicación")
        return False
    
    doc = Document(str(doc_path))
    
    print(f"✓ Documento cargado: {doc_path.name}")
    print(f"  Párrafos actuales: {len(doc.paragraphs)}")
    
    # Buscar puntos de inserción por texto
    graphics_to_insert = [
        {
            'search_text': '5.2.5',
            'graphic_file': 'outputs/ANALISIS_GRAFICO_PVBESSCAR_v7.2.png',
            'caption': 'Figura 5.1: Análisis Gráfico Integral PVBESSCAR\n(Comparación de agentes, sensibilidad, robustez, escalabilidad, Pareto y desglose CO₂)',
            'width': Inches(6.5)
        },
        {
            'search_text': '5.4.1',
            'graphic_file': 'outputs/MATRIZ_SENSIBILIDAD_PESOS.png',
            'caption': 'Figura 5.2: Matriz de Sensibilidad - Impacto de Cambios en Pesos de Recompensa',
            'width': Inches(5.5)
        },
        {
            'search_text': '5.3.5',
            'graphic_file': 'outputs/VALIDACION_TEMPORAL_7DIAS.png',
            'caption': 'Figura 5.3: Validación Temporal - Operación Real (7 días: PV, Demanda, Estado BESS)',
            'width': Inches(6.0)
        }
    ]
    
    inserted_count = 0
    
    for graphic_info in graphics_to_insert:
        graphic_path = Path(graphic_info['graphic_file'])
        
        if not graphic_path.exists():
            print(f"⚠ Gráfico no encontrado: {graphic_path}")
            continue
        
        search_text = graphic_info['search_text']
        
        # Buscar párrafo con el texto de búsqueda
        insertion_index = None
        for i, paragraph in enumerate(doc.paragraphs):
            if search_text in paragraph.text:
                insertion_index = i
                break
        
        if insertion_index is None:
            print(f"⚠ Punto de inserción no encontrado para: {search_text}")
            continue
        
        # Insertar gráfico después del párrafo encontrado
        target_para = doc.paragraphs[insertion_index]
        
        # Crear nuevo párrafo para el gráfico
        new_para = target_para.insert_paragraph_before("")
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Insertar imagen
        try:
            run = new_para.add_run()
            run.add_picture(str(graphic_path), width=graphic_info['width'])
            print(f"✓ Gráfico insertado: {graphic_path.name} ({graphic_info['width'].inches:.1f}\")")
            
            # Agregar caption
            caption_para = target_para.insert_paragraph_before(f"\n{graphic_info['caption']}")
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.runs[0].font.size = Pt(9)
            caption_para.runs[0].font.italic = True
            caption_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            
            inserted_count += 1
            
        except Exception as e:
            print(f"✗ Error insertando gráfico {graphic_path.name}: {str(e)}")
            continue
    
    # Guardar documento con gráficos
    output_path = Path('outputs/TESIS_SECCIONES_5_2_a_5_5_CON_GRAFICOS.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    file_size = output_path.stat().st_size / 1024
    print(f"\n✓ Documento con gráficos guardado: {output_path.name} ({file_size:.1f} KB)")
    print(f"  Total de gráficos insertados: {inserted_count}/3")
    
    return True

if __name__ == '__main__':
    try:
        add_graphics_to_document()
    except Exception as e:
        print(f"\n✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
