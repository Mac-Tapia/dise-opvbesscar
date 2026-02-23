from docx import Document

doc = Document('reports/ARQUITECTURA_PVBESSCAR_TESIS.docx')
print("=" * 70)
print("DOCUMENTO ACTUALIZADO - ÍNDICE DE CONTENIDO")
print("=" * 70)

headings = []
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        level = int(p.style.name.split('Heading')[-1]) if 'Heading' in p.style.name else 0
        indent = "  " * (level - 1) if level > 1 else ""
        headings.append(f"{indent}{p.text[:75]}")

for h in headings:
    print(h)

print(f"\n{'=' * 70}")
print(f"ESTADÍSTICAS:")
print(f"{'=' * 70}")
print(f"✓ Total párrafos: {len([p for p in doc.paragraphs if p.text.strip()])}")
print(f"✓ Tablas: {len(doc.tables)}")
print(f"✓ Encabezados: {len(headings)}")
print(f"✓ Tamaño: {len(open('reports/ARQUITECTURA_PVBESSCAR_TESIS.docx', 'rb').read()) / 1024:.1f} KB")

print(f"\n{'=' * 70}")
print("SECCIONES CLAVE NUEVAS:")
print(f"{'=' * 70}")
for i, p in enumerate(doc.paragraphs):
    text = p.text.lower()
    if any(keyword in text for keyword in ['frecuencia', 'actualización', 'n_steps', 'rollout', 'off-policy', 'continua']):
        if p.style.name.startswith('Heading'):
            print(f"✓ {p.text}")
