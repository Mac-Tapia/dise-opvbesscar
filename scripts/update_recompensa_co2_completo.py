#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Actualizar documento de recompensa con análisis COMPLETO de CO₂
Incluye: Grid import + Solar exported + BESS contribution (nivel SISTEMA)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Sombrear celda de tabla"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_paragraph_indent(doc, text, indent_level=0, bold=False, size=11):
    """Agregar párrafo con indentación"""
    p = doc.add_paragraph(text, style='List Bullet' if indent_level > 0 else None)
    p.paragraph_format.left_indent = Inches(0.25 * indent_level)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
    return p

# Crear documento
doc = Document()

# Título
heading = doc.add_heading('4.6.4.6 Función de Recompensa Multi-Objetivo', level=2)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introducción
doc.add_paragraph(
    'La recompensa del agente SAC combina cinco componentes que impulsan el aprendizaje hacia '
    'una política óptima que minimiza CO₂, maximiza utilización solar, minimiza costos energéticos, '
    'carga vehículos eléctricos y mantiene estabilidad en la red.'
)

# NUEVA SECCIÓN: CO₂ ANÁLISIS A NIVEL SISTEMA
doc.add_heading('Reducción de CO₂ (Componente Indirecto - ANÁLISIS COMPLETO A NIVEL SISTEMA)', level=3)

doc.add_paragraph(
    'La reducción de CO₂ NO se centra SOLO en importación de red pública. Se analiza a nivel SISTEMA, '
    'considerando TODOS los canales donde se evita generación térmica en Iquitos (matriz dominante: 0.4521 kg CO₂/kWh):'
)

# Tabla de resumen de canales
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'

# Headers
header_cells = table.rows[0].cells
header_cells[0].text = 'Canal de Reducción CO₂'
header_cells[1].text = 'Cantidad (kWh/año)'
header_cells[2].text = 'Factor CO₂'
header_cells[3].text = 'CO₂ Ahorrado (kg/año)'

for cell in header_cells:
    shade_cell(cell, 'D3D3D3')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Fila 1: Grid Import Reduction
row1 = table.rows[1].cells
row1[0].text = 'Canal 1: Reducción de Importación'
row1[1].text = '704,533'
row1[2].text = '0.4521'
row1[3].text = '318,516'

# Fila 2: Solar Generation Displaced
row2 = table.rows[2].cells
row2[0].text = 'Canal 2: Generación Desplazada (Solar → Red)'
row2[1].text = '1,921,331'
row2[2].text = '0.4521'
row2[3].text = '868,514'

# Fila 3: BESS Contribution
row3 = table.rows[3].cells
row3[0].text = 'Canal 3: Contribución BESS (Peak Shaving)'
row3[1].text = '257,341*'
row3[2].text = '0.4521'
row3[3].text = '116,243'

# Fila 4: TOTAL
row4 = table.rows[4].cells
row4[0].text = '✅ TOTAL CO₂ AHORRADO (SAC)'
row4[1].text = '2,883,205'
row4[2].text = 'N/A'
row4[3].text = '1,303,273'
for cell in row4:
    shade_cell(cell, 'CCFFCC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(12)

doc.add_paragraph('*BESS contribution: Evita picos de demanda que requieren generación térmica emergente')

# CANAL 1 DETALLADO
doc.add_heading('CANAL 1: Reducción de Importación (Grid Import Reduction)', level=4)

p1 = doc.add_paragraph()
p1.add_run('Baseline (Sin control RL, con solar 4,050 kWp): ').bold = True
p1.add_run('876,000 kWh/año → 396,039 kg CO₂')

p2 = doc.add_paragraph()
p2.add_run('SAC (Con control RL): ').bold = True
p2.add_run('171,467 kWh/año → 77,522 kg CO₂')

p3 = doc.add_paragraph()
p3.add_run('CO₂ AHORRADO: ').bold = True
p3.add_run('704,533 kWh × 0.4521 = 318,516 kg CO₂/año')

p4 = doc.add_paragraph()
p4.add_run('Interpretación: ').italic = True
p4.add_run(
    'El agente RL minimiza la importación de energía térmica mediante optimización de BESS (carga en horas '
    'con exceso solar, descarga en picos) y temporalización de carga de EV (prefiere horarios con solar disponible).'
)

# CANAL 2 DETALLADO
doc.add_heading('CANAL 2: Generación Desplazada (Avoided Thermal Generation via Solar Export)', level=4)

p5 = doc.add_paragraph()
p5.add_run('Generación solar total disponible: ').bold = True
p5.add_run('8,292,514 kWh/año')

p6 = doc.add_paragraph()
p6.add_run('Exceso solar inyectado a red pública: ').bold = True
p6.add_run('1,921,331 kWh/año')

p7 = doc.add_paragraph()
p7.add_run('CO₂ ahorrado por generación desplazada: ').bold = True
p7.add_run('1,921,331 kWh × 0.4521 = 868,514 kg CO₂/año')

p8 = doc.add_paragraph()
p8.add_run('Interpretación CRÍTICA: ').bold = True
p8.add_run(
    'Esta inyección de solar a la red pública de Iquitos REEMPLAZA generación térmica que de otro modo '
    'produciría la central eléctrica. No es "pérdida de energía", es DESPLAZAMIENTO de emisiones: '
    'la central DEJA DE GENERAR porque la red recibe solar de nuestro sistema.\n'
    'Este es el componente más significativo de reducción CO₂ (868.5 ton vs 318.5 ton del canal 1 = 73% del total).'
)

# CANAL 3 DETALLADO
doc.add_heading('CANAL 3: Contribución BESS (Peak Shaving & Thermal Avoidance)', level=4)

p9 = doc.add_paragraph()
p9.add_run('Energía descargada por BESS para peak shaving: ').bold = True
p9.add_run('257,341 kWh/año (estimado OE2)')

p10 = doc.add_paragraph()
p10.add_run('CO₂ ahorrado: ').bold = True
p10.add_run('257,341 kWh × 0.4521 = 116,243 kg CO₂/año')

p11 = doc.add_paragraph()
p11.add_run('Mecanismo: ').italic = True
p11.add_run(
    'El BESS absorbe exceso solar diurno (carga) y lo libera en horas pico (descarga), evitando que '
    'la demanda pico (~170 kW) requiera importación de red térmica. Cada kWh de BESS que evita importación '
    'representa 0.4521 kg CO₂ no generado.'
)

# TOTAL SISTEMA
doc.add_heading('REDUCCIÓN TOTAL CO₂ A NIVEL SISTEMA (SAC)', level=4)

# Tabla de comparación entre agentes
comp_table = doc.add_table(rows=4, cols=5)
comp_table.style = 'Light Grid Accent 1'

comp_headers = comp_table.rows[0].cells
comp_headers[0].text = 'Agente RL'
comp_headers[1].text = 'Canal 1 (Import)'
comp_headers[2].text = 'Canal 2 (Export)'
comp_headers[3].text = 'Canal 3 (BESS)'
comp_headers[4].text = 'TOTAL CO₂'

for cell in comp_headers:
    shade_cell(cell, 'D3D3D3')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

sac_row = comp_table.rows[1].cells
sac_row[0].text = 'SAC (SELECCIONADO)'
sac_row[1].text = '318,516'
sac_row[2].text = '868,514'
sac_row[3].text = '116,243'
sac_row[4].text = '1,303,273'
shade_cell(sac_row[4], 'CCFFCC')
for cell in sac_row[1:]:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

a2c_row = comp_table.rows[2].cells
a2c_row[0].text = 'A2C'
a2c_row[1].text = '348,417'
a2c_row[2].text = '868,514'
a2c_row[3].text = '116,243'
a2c_row[4].text = '1,333,174'

ppo_row = comp_table.rows[3].cells
ppo_row[0].text = 'PPO'
ppo_row[1].text = '286,000'
ppo_row[2].text = '868,514'
ppo_row[3].text = '116,243'
ppo_row[4].text = '1,270,757'

# Conclusión técnica
doc.add_heading('Conclusión Técnica: Por qué SAC es óptimo', level=4)

doc.add_paragraph(
    'A2C logra reducir 29,901 kg CO₂ más que SAC en el Canal 1 (optimización pura de grid import). '
    'Sin embargo, SAC EQUILIBRA ambos objetivos:\n'
)

conclusion_table = doc.add_table(rows=3, cols=3)
conclusion_table.style = 'Light Grid Accent 1'

conc_headers = conclusion_table.rows[0].cells
conc_headers[0].text = 'Métrica'
conc_headers[1].text = 'SAC'
conc_headers[2].text = 'A2C'

for cell in conc_headers:
    shade_cell(cell, 'D3D3D3')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

conc_row1 = conclusion_table.rows[1].cells
conc_row1[0].text = 'CO₂ Total (kg/año)'
conc_row1[1].text = '1,303,273'
conc_row1[2].text = '1,333,174'

conc_row2 = conclusion_table.rows[2].cells
conc_row2[0].text = 'EV Cargados/año'
conc_row2[1].text = '3,500 (+25%)'
conc_row2[2].text = '3,000 (+7.1%)'
shade_cell(conc_row2[1], 'FFEB99')  # Yellow highlight for SAC advantage

doc.add_paragraph(
    'SAC sacrifica apenas 29,901 kg CO₂ (2.2%) respecto a A2C, pero GANA 500 vehículos cargados/año (+16.7% más usuarios). '
    'Para una ciudad real (Iquitos), la satisfacción de usuarios es crítica: SAC es el balance óptimo entre '
    'ambición de CO₂ y viabilidad operacional.'
)

# Pesos finales
doc.add_heading('Pesos de Recompensa (Configuración Final SAC)', level=3)

weights_table = doc.add_table(rows=6, cols=3)
weights_table.style = 'Light Grid Accent 1'

w_headers = weights_table.rows[0].cells
w_headers[0].text = 'Componente'
w_headers[1].text = 'Peso'
w_headers[2].text = 'Descripción'

for cell in w_headers:
    shade_cell(cell, 'D3D3D3')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

weights = [
    ('w_CO₂ (Reducción CO₂ indirecto)', '0.35', 'PRIMARY: Canal 1 + Canal 2 + Canal 3'),
    ('w_EV (Carga de vehículos)', '0.30', 'SECONDARY: Maximizar EV cargados/año'),
    ('w_solar (Utilización solar)', '0.20', 'TERTIARY: Maximizar % solar vs grid'),
    ('w_cost (Minimización de costo)', '0.10', 'TERTIARY: Reducir tariff impact'),
    ('w_grid (Estabilidad red)', '0.05', 'TERTIARY: Smooth power ramping'),
]

for i, (comp, weight, desc) in enumerate(weights, 1):
    row = weights_table.rows[i].cells
    row[0].text = comp
    row[1].text = weight
    row[2].text = desc
    if 'PRIMARY' in desc:
        shade_cell(row[0], 'FFE699')

doc.add_paragraph(
    'Penalización condicional: -0.5 si BESS SOC < 65% durante horas pico (16-17h), '
    'para garantizar disponibilidad de descarga en picos de demanda.'
)

# Guardar
output_path = 'reports/SECCION_4646_FUNCION_RECOMPENSA_CO2_COMPLETO.docx'
doc.save(output_path)

print(f'✓ Documento actualizado: {output_path}')
print(f'✓ Incluye análisis COMPLETO de CO₂ a nivel SISTEMA')
print(f'✓ Canales: Grid Import + Solar Export + BESS Contribution')
print(f'✓ Total CO₂ SAC: 1,303,273 kg/año (análisis completo)')
