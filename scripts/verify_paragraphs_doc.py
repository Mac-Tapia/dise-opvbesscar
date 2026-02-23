from docx import Document

doc = Document('reports/PARRAFOS_LISTOS_PARA_PEGAR.docx')

print('='*70)
print('DOCUMENTO: PÁRRAFOS LISTOS PARA PEGAR')
print('='*70)
size = len(open('reports/PARRAFOS_LISTOS_PARA_PEGAR.docx', 'rb').read()) / 1024
paras = len([p for p in doc.paragraphs if p.text.strip()])
tables = len(doc.tables)
print(f'✓ Tamaño: {size:.1f} KB')
print(f'✓ Párrafos: {paras}')
print(f'✓ Tablas: {tables}')

print()
print('ESTRUCTURA:')
print('='*70)
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        level = int(p.style.name.split('Heading')[-1]) if 'Heading' in p.style.name else 0
        indent = '  ' * (level - 1) if level > 1 else ''
        text = p.text[:65]
        print(f'{indent}• {text}')

print()
print('='*70)
print('✅ LISTO PARA USAR EN: reports/PARRAFOS_LISTOS_PARA_PEGAR.docx')
print('='*70)
