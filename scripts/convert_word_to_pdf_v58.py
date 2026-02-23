#!/usr/bin/env python3
"""
Script para convertir Word v5.8 a PDF
Usa LibreOffice o MS Word COM si está disponible
"""

from pathlib import Path
import subprocess
import os
import sys

DOC_SOURCE = Path("outputs/docx/BESS_Dimensionamiento_Procedimiento_v5.8_2026-02-21.docx")
PDF_OUTPUT = Path("outputs/pdf/BESS_Dimensionamiento_v5.8_2026-02-21.pdf")

print("=" * 80)
print("CONVERSIÓN: Word v5.8 → PDF v5.8")
print("=" * 80)

if not DOC_SOURCE.exists():
    print(f"❌ ERROR: No existe {DOC_SOURCE}")
    sys.exit(1)

print(f"📄 Documento origen: {DOC_SOURCE}")
print(f"📄 PDF destino: {PDF_OUTPUT}")

# Crear directorio si no existe
PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)

# Opción 1: Intentar con LibreOffice
print("\n🔍 Intentando con LibreOffice...")
try:
    # En Windows
    libreoffice_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    
    lo_path = None
    for path in libreoffice_paths:
        if os.path.exists(path):
            lo_path = path
            break
    
    if lo_path:
        print(f"✅ LibreOffice encontrado en: {lo_path}")
        cmd = [
            lo_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(PDF_OUTPUT.parent),
            str(DOC_SOURCE)
        ]
        print(f"🔄 Ejecutando: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, timeout=60)
        
        # Verificar si se creó el PDF
        expected_pdf = PDF_OUTPUT.parent / f"{DOC_SOURCE.stem}.pdf"
        if expected_pdf.exists():
            # Renombrar a nombre deseado si es diferente
            if expected_pdf != PDF_OUTPUT:
                expected_pdf.rename(PDF_OUTPUT)
            print(f"✅ PDF creado con LibreOffice: {PDF_OUTPUT}")
            print("\n" + "=" * 80)
            print("✅ CONVERSIÓN COMPLETADA!")
            print("=" * 80)
            sys.exit(0)
        else:
            print(f"⚠️ LibreOffice no generó PDF")
    else:
        print(f"⚠️ LibreOffice no encontrado en rutas estándar")
except Exception as e:
    print(f"⚠️ Error con LibreOffice: {e}")

# Opción 2: Intentar con Word COM (Windows only)
print("\n🔍 Intentando con MS Word COM...")
try:
    from win32com.client import Dispatch
    
    print("✅ win32com disponible")
    word = Dispatch("Word.Application")
    word.Visible = False
    
    doc = word.Documents.Open(str(DOC_SOURCE.absolute()))
    
    # Convertir a PDF
    doc.SaveAs(
        str(PDF_OUTPUT.absolute()),
        FileFormat=17  # wdFormatPDF = 17
    )
    
    doc.Close()
    word.Quit()
    
    print(f"✅ PDF creado con MS Word: {PDF_OUTPUT}")
    print("\n" + "=" * 80)
    print("✅ CONVERSIÓN COMPLETADA!")
    print("=" * 80)
    sys.exit(0)
    
except Exception as e:
    print(f"⚠️ Error con MS Word COM: {e}")

# Opción 3: Python-pptx (no funcionará para docx, pero es alternativa)
print("\n🔍 Intentando con python-docx + reportlab...")
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table
    from reportlab.lib.styles import getSampleStyleSheet
    
    print("✅ Librerías necesarias disponibles")
    
    # Leer documento
    doc = Document(str(DOC_SOURCE))
    
    # Crear PDF
    pdf_doc = SimpleDocTemplate(str(PDF_OUTPUT), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            p = Paragraph(para.text, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.2))
    
    # Agregar tablas
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        if table_data:
            t = Table(table_data)
            story.append(t)
            story.append(Spacer(1, 0.3))
    
    pdf_doc.build(story)
    
    print(f"✅ PDF creado con reportlab: {PDF_OUTPUT}")
    print("\n" + "=" * 80)
    print("✅ CONVERSIÓN COMPLETADA!")
    print("=" * 80)
    sys.exit(0)
    
except Exception as e:
    print(f"⚠️ Error con reportlab: {e}")

# Si llegamos aquí, ningún método funcionó
print("\n" + "=" * 80)
print("❌ NO se pudo convertir a PDF automáticamente")
print("\nAlternativas manuales:")
print("1. Abrir documento en MS Word")
print(f"   Archivo: {DOC_SOURCE}")
print("2. File → Save As")
print("3. Cambiar formato a PDF")
print(f"4. Guardar como: {PDF_OUTPUT}")
print("=" * 80)
sys.exit(1)
