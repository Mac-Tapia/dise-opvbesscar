#!/usr/bin/env python3
"""
Crear documento integrado combinando secciones 5.2-5.3-5.4-5.5
Genera un Word único con tabla de contenidos y referencias cruzadas
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Agregar src al path
src_path = Path(__file__).parent.parent
sys.path.insert(0, str(src_path))

def create_integrated_document():
    """Combinar secciones 5.2-5.5 en un único documento"""
    
    print("\n" + "="*80)
    print("CREANDO DOCUMENTO INTEGRADO 5.2-5.5")
    print("="*80 + "\n")
    
    # Crear documento maestro
    integrated_doc = Document()
    
    # Agregar portada/encabezado
    title = integrated_doc.add_heading(
        'SECCIONES 5.2 - 5.5\nDimensionamiento, Algoritmo RL, Sensibilidad y Conclusiones',
        level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = integrated_doc.add_paragraph(
        'Sistema de Optimización de Carga EV con Energía Solar y BESS\nPVBESSCAR v7.2'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].italic = True
    
    integrated_doc.add_paragraph()  # Espacio
    
    # Tabla de contenidos (simplificada)
    toc_heading = integrated_doc.add_heading('TABLA DE CONTENIDOS', level=1)
    toc_items = [
        ('5.2', 'Dimensionamiento de Infraestructura'),
        ('5.2.1', 'Capacidad de Generación Solar'),
        ('5.2.2', 'Dimensionamiento de Cargadores EV'),
        ('5.2.3', 'Dimensionamiento del BESS'),
        ('5.2.4', 'Balance Energético Diario'),
        ('5.2.5', 'Contribución CO₂ por Infraestructura'),
        ('5.3', 'Algoritmo de Control RL'),
        ('5.3.1', 'Estrategia SAC (Soft Actor-Critic)'),
        ('5.3.2', 'Función de Recompensa Multi-Objetivo'),
        ('5.3.3', 'Resultados Comparativos'),
        ('5.3.4', 'Mecanismo de Optimización'),
        ('5.3.5', 'Validación contra Datos Reales'),
        ('5.4', 'Análisis de Sensibilidad y Consideraciones'),
        ('5.4.1', 'Sensibilidad a Pesos de Recompensa'),
        ('5.4.2', 'Robustez ante Perturbaciones'),
        ('5.4.3', 'Escalabilidad'),
        ('5.4.4', 'Consideraciones Operacionales'),
        ('5.5', 'Validación de Hipótesis y Conclusiones'),
        ('5.5.1', 'Validación de Hipótesis'),
        ('5.5.2', 'Contribución Científica'),
        ('5.5.3', 'Conclusiones Generales'),
        ('5.5.4', 'Recomendaciones de Implementación'),
        ('5.5.5', 'Síntesis Final'),
    ]
    
    for section, title in toc_items:
        p = integrated_doc.add_paragraph(f'{section}: {title}', style='List Number')
        p.paragraph_format.left_indent = Inches(0.5 * int(section.count('.')))
    
    integrated_doc.add_page_break()
    
    # Cargar y copiar contenido de cada sección
    output_dir = Path('outputs')
    
    # Archivos de secciones a integrar
    section_files = [
        'SECCION_5_2_DIMENSIONAMIENTO_DESCRIPTIVO_COMPLETO.docx',
        'SECCION_5_3_ALGORITMO_RL_COMPLETO.docx',
        'SECCION_5_4_SENSIBILIDAD_OPERACIONAL_COMPLETA.docx',
        'SECCION_5_5_VALIDACION_CONCLUSIONES_FINALES.docx'
    ]
    
    for section_file in section_files:
        section_path = output_dir / section_file
        if not section_path.exists():
            print(f"⚠ Archivo no encontrado: {section_file}")
            continue
        
        print(f"✓ Leyendo contenido de {section_file}...")
        
        try:
            section_doc = Document(str(section_path))
            
            # Copiar párrafos y tablas
            for element in section_doc.element.body:
                # Copiar elemento al documento integrado
                integrated_doc.element.body.append(element.__deepcopy__(None))
            
            integrated_doc.add_page_break()
            
        except Exception as e:
            print(f"✗ Error leyendo {section_file}: {str(e)}")
            continue
    
    # Guardar documento integrado
    output_file = 'TESIS_SECCIONES_5_2_a_5_5_INTEGRADO_COMPLETO.docx'
    integrated_doc.save(output_file)
    
    file_size = os.path.getsize(output_file) / 1024  # KB
    print(f"\n✓ Documento integrado generado: {output_file} ({file_size:.1f} KB)")
    
    # Resumen de contenido
    print("\n" + "="*80)
    print("CONTENIDO INTEGRADO:")
    print("="*80)
    print("""
Sección 5.2: Dimensionamiento Descriptivo
  → 5.2.1: Generación Solar (2 tablas: patrones horarios + mensuales)
  → 5.2.2: Cargadores EV (1 tabla: especificaciones 19 chargers × 2 sockets)
  → 5.2.3: BESS (1 tabla: 10 parámetros de batería)
  → 5.2.4: Balance Energético (1 tabla: ciclo diario 6 períodos)
  → 5.2.5: Contribución CO₂ (1 tabla: 3 canales)
  
Sección 5.3: Algoritmo RL
  → 5.3.1: Estrategia SAC (1 tabla: parámetros de entrenamiento)
  → 5.3.2: Función Recompensa (1 tabla: 5 componentes de reward)
  → 5.3.3: Resultados Comparativos (1 tabla: SAC vs PPO vs A2C)
  → 5.3.4: Mecanismo de Optimización (descripción)
  → 5.3.5: Validación (descripción)
  
Sección 5.4: Análisis de Sensibilidad
  → 5.4.1: Sensibilidad a Pesos (1 tabla: 5 escenarios)
  → 5.4.2: Robustez ante Perturbaciones (1 tabla: 6 escenarios)
  → 5.4.3: Escalabilidad (1 tabla: 5 proyecciones)
  → 5.4.4: Consideraciones Operacionales (descripción + 5 items)
  
Sección 5.5: Validación y Conclusiones
  → 5.5.1: Validación Hipótesis (1 tabla: 6 hipótesis)
  → 5.5.2: Contribución Científica (4 puntos)
  → 5.5.3: Conclusiones Generales (6 conclusiones)
  → 5.5.4: Recomendaciones Implementación (3 fases)
  → 5.5.5: Síntesis Final (viabilidad + números)
    """)
    
    # Estadísticas del documento
    print("\n" + "="*80)
    print("ESTADÍSTICAS DEL DOCUMENTO INTEGRADO:")
    print("="*80)
    print(f"Total de párrafos: {len(integrated_doc.paragraphs)}")
    print(f"Total de tablas: {len(integrated_doc.tables)}")
    print(f"Tamaño del archivo: {file_size:.1f} KB")
    
    print("\n✓ Documento integrado completado exitosamente")
    return True

if __name__ == '__main__':
    try:
        create_integrated_document()
    except Exception as e:
        print(f"\n✗ Error: {str(e)}")
        sys.exit(1)
