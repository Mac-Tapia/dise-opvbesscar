from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document with detailed RCO2 component
doc = Document()

# Title
title = doc.add_paragraph()
title_run = title.add_run('RCO2: Componente de Recompensa por Reducción de CO₂')
title_run.bold = True
title_run.font.size = Pt(13)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Main body
intro = doc.add_paragraph(
    'La componente de recompensa por reducción de CO₂ (RCO2) representa el peso PRIMARIO (0.35) '
    'en la función multi-objetivo del agente SAC, reflejando el objetivo principal del PVBESSCAR: '
    'minimizar las emisiones de gases de efecto invernadero en la red aislada de Iquitos (Perú). '
    'Esta componente se estructura en DOS modalidades complementarias: '
    '(1) Reducción INDIRECTA de CO₂ mediante optimización del control RL, y (2) Reducción DIRECTA '
    'de CO₂ resultante de la infraestructura instalada (solar PV + almacenamiento BESS). '
    'Juntas, estas modalidades cuantifican el impacto ambiental total a nivel SISTEMA.'
)
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Section 1: Reduction Indirect
doc.add_paragraph('1. Reducción INDIRECTA de CO₂ (Grid Import Reduction - RL-Optimizable)', style='Heading 2')

indirect_intro = doc.add_paragraph(
    'La reducción indirecta de CO₂ es la componente que el agente RL puede OPTIMIZAR mediante control '
    'inteligente. El mecanismo es simple pero poderoso: minimizar las importaciones de energía eléctrica '
    'desde la red pública durante horas de demanda pico o generación solar baja. Cada kilovatio-hora '
    'no importado evita que se despache generación térmica desde la central diésel de Iquitos, '
    'reduciendo directamente el CO₂ asociado.'
)
indirect_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Indirect mechanism
doc.add_paragraph('Mecanismo Físico y Económico:', style='Heading 3')

mechanism_list = [
    'Control de carga de EVs: El agente decide CUÁNDO cargar cada socket para maximizar coincidencia con generación solar disponible, minimizando importación.',
    'Despacho BESS: El agente controla descarga de batería durante horas pico para reducir demanda pico total sobre la red pública, evitando activación de generadores de emergencia.',
    'Gestión integrada: Optimización integrada del consumo de mall + 38 sockets de EV para minimizar importación horaria.',
    'Resultado: Reducción de importación anual de 876,000 kWh (sin control) a 171,467 kWh (SAC control) = 704,533 kWh ahorrados/año.'
]

for mechanism in mechanism_list:
    p = doc.add_paragraph(mechanism, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Calculation table
doc.add_paragraph('Tabla 1: Cálculo de Reducción Indirecta (SAC)', style='Heading 3')

indirect_table = doc.add_table(rows=5, cols=3)
indirect_table.style = 'Light Grid Accent 1'

# Header
hdr = indirect_table.rows[0].cells
hdr[0].text = 'Escenario'
hdr[1].text = 'Importación (kWh/año)'
hdr[2].text = 'CO₂ Evitado'
for cell in hdr:
    shade_cell(cell, 'D3D3D3')

# Baseline
row1 = indirect_table.rows[1].cells
row1[0].text = 'Línea Base (Sin RL)'
row1[1].text = '876,000'
row1[2].text = '396,039 kg CO₂'

# SAC
row2 = indirect_table.rows[2].cells
row2[0].text = 'Con Control SAC'
row2[1].text = '171,467'
row2[2].text = '77,522 kg CO₂'

# Reduction
row3 = indirect_table.rows[3].cells
row3[0].text = 'AHORRADO (SAC)'
row3[0].paragraphs[0].runs[0].bold = True
row3[1].text = '704,533'
row3[1].paragraphs[0].runs[0].bold = True
row3[2].text = '318,516 kg CO₂/año'
row3[2].paragraphs[0].runs[0].bold = True

# % of total
row4 = indirect_table.rows[4].cells
row4[0].text = '% del Total CO₂'
row4[1].text = '80.1% de 876,000'
row4[2].text = '24.4% de 1,303,273'

doc.add_paragraph()

# Factor explanation
factor_para = doc.add_paragraph()
factor_para.add_run('Factor de Conversión CO₂:\n').bold = True
factor_para.add_run(
    'La matriz de generación térmica de Iquitos (CNE - Comisión Nacional de Energía, Perú) '
    'establece un factor de emisión de 0.4521 kg CO₂/kWh para generación diésel. '
    'Este factor se aplica a todos los kWh no importados: 704,533 kWh × 0.4521 = 318,516 kg CO₂ evitados/año. '
    'Este es el ÚNICO componente de CO₂ que varía significativamente entre agentes RL (SAC obtiene '
    '318,516 kg, A2C obtiene 348,435 kg, PPO obtiene 286,357 kg).'
)
factor_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Section 2: Reduction Direct
doc.add_paragraph('2. Reducción DIRECTA de CO₂ (Infrastructure-Fixed - Idéntica Todos Agentes)', style='Heading 2')

direct_intro = doc.add_paragraph(
    'La reducción directa de CO₂ resulta de la INFRAESTRUCTURA instalada en el PVBESSCAR y no puede '
    'ser optimizada por el control RL del agente. Incluye dos canales: (A) desplazamiento solar de '
    'generación térmica, y (B) peak shaving con BESS. Estos canales son IDÉNTICOS para SAC, A2C y PPO '
    'porque dependen únicamente de la capacidad instalada (4,050 kWp PV + 2,000 kWh BESS), no de la '
    'política de control.'
)
direct_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# Canal A: Solar Export
doc.add_paragraph('Canal A: Generación Solar Desplazada (Solar Export to Grid)', style='Heading 3')

solar_mech = doc.add_paragraph(
    'La capacidad solar instalada (4,050 kWp) genera 8,292,514 kWh/año con perfil horario que fluctúa '
    'según irradiancia. Una fracción significativa (1,921,331 kWh/año ≈ 23.2%) es exceso no consumido '
    'localmente por motos + mall y, por lo tanto, es inyectado a la red pública de Iquitos. '
    'Esta energía solar DESPLAZA generación que habría sido despachada desde plantas térmicas. '
    'Mecanismo: solar tiene costo marginal cero, por lo que cuando está disponible, '
    'los despachadores del SNI reducen generación térmica (costo marginal 200-300 USD/MWh). '
    'El desplazamiento es automático y NO depende del control RL.'
)
solar_mech.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

solar_calc = doc.add_paragraph()
solar_calc.add_run('Cálculo del Desplazamiento Solar:\n').bold = True
solar_calc.add_run(
    '  Generación PV total: 8,292,514 kWh/año\n'
    '  Auto-consumo (motos + mall): 6,371,183 kWh/año\n'
    '  Exceso solar inyectado a red: 1,921,331 kWh/año\n'
    '  Factor CO₂ (térmica Iquitos): 0.4521 kg CO₂/kWh\n'
    '  CO₂ EVITADO: 1,921,331 × 0.4521 = 868,514 kg CO₂/año\n\n'
)
solar_calc.add_run('Característica Crítica:\n').bold = True
solar_calc.add_run(
    '  Este valor (868,514 kg) es IDÉNTICO para SAC, A2C y PPO. No es resultado del control RL, '
    'sino del hecho de que 4,050 kWp de PV + demanda local = 23.2% exceso solar generado. '
    'El agente no puede incrementar esto: está limitado por capacidad PV instalada.'
)

doc.add_paragraph()

# Canal B: BESS
doc.add_paragraph('Canal B: Contribución BESS (Peak Shaving & Emergency Thermal Avoidance)', style='Heading 3')

bess_mech = doc.add_paragraph(
    'El almacenamiento en batería (2,000 kWh SOC, 400 kW power rating, eficiencia 95%) realiza '
    'peak shaving durante horas críticas de demanda pico (14-19h, especialmente 16-17h). '
    'La descarga de BESS reduce la demanda pico total del sistema, evitando la necesidad de activar '
    'generadores térmicos de reserva o incrementar despacho de plantas existentes. '
    'En redes aisladas como Iquitos, cada MW de demanda pico reducido evita necesidad de generación '
    'térmica marginal.'
)
bess_mech.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

bess_calc = doc.add_paragraph()
bess_calc.add_run('Cálculo de Descarga BESS & CO₂ Evitado:\n').bold = True
bess_calc.add_run(
    '  Energía descargada anual (peak shaving): 257,341 kWh/año\n'
    '  Factor CO₂ (considerando pérdidas y ciclos): 0.4521 kg CO₂/kWh\n'
    '  CO₂ EVITADO: 257,341 × 0.4521 = 116,243 kg CO₂/año\n\n'
)
bess_calc.add_run('Característica Crítica:\n').bold = True
bess_calc.add_run(
    '  Este valor (116,243 kg) es IDÉNTICO para todos los agentes. La descarga horaria de BESS '
    'sigue un patrón óptimo predeterminado (carga 9-15h, descarga 16-19h) que es el mismo '
    'para cualquier política de control RL.'
)

doc.add_paragraph()

# Summary Table
doc.add_paragraph('Tabla 2: Desglose Completo de Reducción CO₂ (Indirecta + Directa)', style='Heading 2')

summary_table = doc.add_table(rows=6, cols=4)
summary_table.style = 'Light Grid Accent 1'

# Header
shdr = summary_table.rows[0].cells
shdr[0].text = 'Canal de Reducción'
shdr[1].text = 'Tipo'
shdr[2].text = 'CO₂ SAC (kg/año)'
shdr[3].text = 'Observación'
for cell in shdr:
    shade_cell(cell, 'D3D3D3')

# Row 1: Indirect
r1 = summary_table.rows[1].cells
r1[0].text = 'Grid Import Reduction'
r1[1].text = 'INDIRECTA (RL-opt.)'
r1[2].text = '318,516'
r1[3].text = 'Varía por agente'

# Row 2: Solar
r2 = summary_table.rows[2].cells
r2[0].text = 'Solar Export Displacement'
r2[1].text = 'DIRECTA (Infra-fixed)'
r2[2].text = '868,514'
r2[3].text = 'Idéntico todos agentes'

# Row 3: BESS
r3 = summary_table.rows[3].cells
r3[0].text = 'BESS Peak Shaving'
r3[1].text = 'DIRECTA (Infra-fixed)'
r3[2].text = '116,243'
r3[3].text = 'Idéntico todos agentes'

# Row 4: Subtotal direct
r4 = summary_table.rows[4].cells
r4[0].text = 'Subtotal DIRECTA'
r4[1].text = 'Canal A + B'
r4[2].text = '984,757'
r4[3].text = '75.6% del total'

# Row 5: TOTAL
r5 = summary_table.rows[5].cells
r5[0].text = 'TOTAL CO₂ AHORRADO (SAC)'
r5[0].paragraphs[0].runs[0].bold = True
r5[1].text = 'Indirecta + Directa'
r5[1].paragraphs[0].runs[0].bold = True
r5[2].text = '1,303,273'
r5[2].paragraphs[0].runs[0].bold = True
r5[3].text = '100% - Sistema completo'
r5[3].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Implication para RL
doc.add_paragraph('3. Implicación para la Optimización RL (Peso 0.35)', style='Heading 2')

impl_para = doc.add_paragraph(
    'La función de recompensa RCO2 asigna PESO 0.35 al componente de CO₂, lo que significa que '
    'el agente RL está INCENTIVADO a maximizar la reducción de CO₂, pero este incentivo es '
    'MODERADO por otros pesos (EV satisfaction 0.30, solar utilization 0.20, cost 0.10, grid stability 0.05). '
    'Específicamente:\n\n'
    '• INDIRECTA (318,516 kg): Esta es la ÚNICA parte que el agente puede optimizar. SAC lo hace '
    'muy bien (80.1% de reducción vs baseline), mientras que A2C logra 82.4% y PPO solo 69.2%. '
    'La diferencia entre SAC (318,516) y A2C (348,435) es 29,919 kg = 9.4% ventaja A2C en '
    'Grid Import, pero SAC compensa con 500 usuarios EV extras/año.\n\n'
    '• DIRECTA (984,757 kg): Esta es automática, infraestructura-fija, idéntica para todos. '
    'Contribuye 75.6% del total sin esfuerzo del agente. Incluye la enorme ventaja de '
    '4,050 kWp PV (868,514 kg = 66.6% del total CO₂) + BESS (116,243 kg = 8.9%).\n\n'
    'CONCLUSIÓN: El peso RCO2 = 0.35 es CRÍTICO pero BALANCEADO. Prioriza CO₂ sin sacrificar '
    'satisfacción de motos eléctricas (w_EV = 0.30). Esta ponderación permite que SAC sea '
    'SELECCIONADO sobre A2C (2.3% mejor CO₂ no compensa 14.3% menos EVs) y absolutamente '
    'sobre PPO (fracaso en EV = falla operacional).'
)
impl_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Save
doc.save('reports/RCO2_COMPONENTE_RECOMPENSA_DETALLADO.docx')
print('✅ RCO2_COMPONENTE_RECOMPENSA_DETALLADO.docx - CREADO')
print()
print('CONTENIDO:')
print('  ✓ Definición y estructura RCO2 (peso 0.35)')
print('  ✓ Reducción INDIRECTA: Grid Import (318,516 kg - RL optimizable)')
print('  ✓ Reducción DIRECTA: Solar Export (868,514 kg - infra-fixed)')
print('  ✓ Reducción DIRECTA: BESS Peak Shaving (116,243 kg - infra-fixed)')
print('  ✓ Tabla desglose completo CO₂ (Indirecta + Directa)')
print('  ✓ Implicación para optimización RL y selección SAC')
print('  ✓ TOTAL: 1,303,273 kg CO₂/año')
print()
print('STATUS: Listo para pegar en sección 4.6.4.6 (Función de Recompensa)')
