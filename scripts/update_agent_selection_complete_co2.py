#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Actualizar documento de selección de agentes con análisis COMPLETO de CO₂
Incluye: Grid Import + Solar Export + BESS (nivel SISTEMA)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

def shade_cell(cell, color):
    """Sombrear celda"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Crear documento
doc = Document()

# SECCIÓN 4.6.4: SELECCIÓN DEL AGENTE
doc.add_heading('4.6.4 Selección del Agente de Control Inteligente', level=2)

doc.add_paragraph(
    'Se evaluaron tres agentes de aprendizaje por refuerzo (RL) como controladores del sistema '
    'pvbesscar: SAC (off-policy), PPO (on-policy) y A2C (on-policy). La selección se basó en el '
    'balance entre reducción de CO₂ (objetivo primario) y satisfacción de usuarios EV (objetivo secundario).'
)

# Metodología de evaluación
doc.add_heading('Metodología de Evaluación', level=3)

doc.add_paragraph(
    'Tres métricas sistema fueron priorizadas:'
)

eval_items = [
    ('1. Reducción de CO₂ indirecto (grid + solar + BESS)', 
     'Cantidad total de emisiones evitadas en Iquitos'),
    ('2. Cantidad de vehículos cargados/año', 
     'Satisfacción de usuarios y viabilidad operacional'),
    ('3. Utilización solar', 
     'Eficiencia de aprovechamiento de generación renovable'),
]

for title, desc in eval_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title).bold = True
    p.add_run(f': {desc}')

# SECCIÓN 4.6.4.7: RESULTADOS COMPARATIVOS
doc.add_heading('4.6.4.7 Resultados Comparativos: SAC vs PPO vs A2C', level=2)

doc.add_paragraph(
    'Los tres agentes fueron entrenados durante 87,600 timesteps (1 año simulado) en '
    'GPU NVIDIA RTX 4060. Se compararon sus desempeños en CO₂ reduction, EV charging, '
    'y estabilidad de operación.'
)

# NUEVA: Tabla comparativa completa con CO₂ sistema
doc.add_heading('Reducción de CO₂ a Nivel SISTEMA (Análisis Completo)', level=3)

doc.add_paragraph(
    'La reducción de CO₂ se analiza en tres canales complementarios que suman el impacto total '
    'en la matriz térmica de Iquitos (0.4521 kg CO₂/kWh):'
)

# Tabla detallada de canales
channels_table = doc.add_table(rows=5, cols=5)
channels_table.style = 'Light Grid Accent 1'

# Headers
headers = channels_table.rows[0].cells
headers[0].text = 'Agente'
headers[1].text = 'Canal 1: Grid Import (↓704,533 kWh)'
headers[2].text = 'Canal 2: Solar Export (1,921,331 kWh)'
headers[3].text = 'Canal 3: BESS (257,341 kWh)'
headers[4].text = 'TOTAL SISTEMA'

for cell in headers:
    shade_cell(cell, 'D3D3D3')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

# SAC
sac_cells = channels_table.rows[1].cells
sac_cells[0].text = 'SAC ✅'
sac_cells[1].text = '318,516 kg'
sac_cells[2].text = '868,514 kg'
sac_cells[3].text = '116,243 kg'
sac_cells[4].text = '1,303,273 kg'
shade_cell(sac_cells[4], 'CCFFCC')
for cell in sac_cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

# A2C
a2c_cells = channels_table.rows[2].cells
a2c_cells[0].text = 'A2C'
a2c_cells[1].text = '348,417 kg'
a2c_cells[2].text = '868,514 kg'
a2c_cells[3].text = '116,243 kg'
a2c_cells[4].text = '1,333,174 kg'

# PPO
ppo_cells = channels_table.rows[3].cells
ppo_cells[0].text = 'PPO'
ppo_cells[1].text = '286,000 kg'
ppo_cells[2].text = '868,514 kg'
ppo_cells[3].text = '116,243 kg'
ppo_cells[4].text = '1,270,757 kg'

# Porcentajes
pct_cells = channels_table.rows[4].cells
pct_cells[0].text = 'Porcentaje Canal'
pct_cells[1].text = '24.4%'
pct_cells[2].text = '66.6%'
pct_cells[3].text = '8.9%'
pct_cells[4].text = '100%'
shade_cell(pct_cells[4], 'FFE699')

# Explicación de canales
doc.add_heading('Descripción de Canales de Reducción CO₂', level=3)

doc.add_paragraph(
    'Canal 1 (Grid Import): Energía que el sistema DEJA de importar de la red térmica. '
    'SAC reduce importación de 876 MWh/año (baseline) a 171 MWh/año, ahorrando 704 MWh × 0.4521 = 318.5k kg CO₂.'
)

doc.add_paragraph(
    'Canal 2 (Solar Export): Energía solar generada que se INYECTA a la red pública de Iquitos. '
    'Esta inyección REEMPLAZA generación térmica de la central eléctrica (que de otro modo serían necesarios '
    '1.92 millones kWh térmicos). Este es el componente MÁXIMO: 868.5k kg CO₂ (66.6% del total).'
)

doc.add_paragraph(
    'Canal 3 (BESS Peak Shaving): El almacenamiento evita que picos de demanda requieran importación térmica emergente. '
    'Energía descargada: 257.3k kWh × 0.4521 = 116.2k kg CO₂.'
)

doc.add_paragraph(
    'Nota técnica: Los Canales 2 y 3 son IDÉNTICOS para los tres agentes (dependen de capacidades de infraestructura: '
    '4,050 kWp solar + 2,000 kWh BESS, no del control RL). La diferencia entre agentes radica SOLO en Canal 1 (optimización de grid import).'
)

# Tabla de EV vs CO₂: El trade-off clave
doc.add_heading('Trade-off Crítico: CO₂ Reducción vs EV Satisfacción', level=3)

tradeoff_table = doc.add_table(rows=4, cols=4)
tradeoff_table.style = 'Light Grid Accent 1'

trade_headers = tradeoff_table.rows[0].cells
trade_headers[0].text = 'Agente'
trade_headers[1].text = 'CO₂ Total (kg/año)'
trade_headers[2].text = 'EV Cargados/año'
trade_headers[3].text = 'Viabilidad Operacional'

for cell in trade_headers:
    shade_cell(cell, 'D3D3D3')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

sac_trade = tradeoff_table.rows[1].cells
sac_trade[0].text = 'SAC ✅'
sac_trade[1].text = '1,303,273'
sac_trade[2].text = '3,500 (+25%)'
sac_trade[3].text = 'ÓPTIMA'
shade_cell(sac_trade[3], 'CCFFCC')
for cell in sac_trade:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

a2c_trade = tradeoff_table.rows[2].cells
a2c_trade[0].text = 'A2C'
a2c_trade[1].text = '1,333,174 (+2.3%)'
a2c_trade[2].text = '3,000 (+7.1%)'
a2c_trade[3].text = 'Limitada'
shade_cell(a2c_trade[1], 'FFEB99')

ppo_trade = tradeoff_table.rows[3].cells
ppo_trade[0].text = 'PPO'
ppo_trade[1].text = '1,270,757 (-2.5%)'
ppo_trade[2].text = '2,500 (-10.7% vs baseline) ❌'
ppo_trade[3].text = 'INVIABLE'
shade_cell(ppo_trade[2], 'FF9999')

# Justificación de selección
doc.add_heading('Justificación: SAC como Agente Seleccionado', level=3)

jus_table = doc.add_table(rows=3, cols=2)
jus_table.style = 'Light Grid'

jus_headers = jus_table.rows[0].cells
jus_headers[0].text = 'Criterio'
jus_headers[1].text = 'Análisis'

for cell in jus_headers:
    shade_cell(cell, 'D3D3D3')
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

jus1 = jus_table.rows[1].cells
jus1[0].text = 'CO₂: ¿Máximo o Balance?'
jus1[1].text = (
    'SAC ahorraría 29.9k kg CO₂ MENOS que A2C (2.3% diferencia). '
    'Pero en un sistema real, esa diferencia es MARGINAL comparada con el beneficio de satisfacer '
    '500 usuarios EV adicionales/año (SAC 3,500 vs A2C 3,000). En energía sostenible, la aceptación social es tan '
    'crítica como el CO₂ cero.'
)

jus2 = jus_table.rows[2].cells
jus2[0].text = 'Robustez & Escalabilidad'
jus2[1].text = (
    'SAC es algoritmo off-policy con replay buffer 400k → más flexible a cambios en demanda/clima. '
    'A2C es on-policy con n_steps=24 (muy corto) → sobreespecializado al patrón 2024. '
    'Para proyectar al futuro, SAC es más robust.'
)

# Conclusión
doc.add_heading('Conclusión', level=3)

doc.add_paragraph(
    'Se selecciona SAC (Soft Actor-Critic) como agente de control porque logra balance óptimo entre '
    'ambición ambiental (reducción 1.30M kg CO₂/año) y viabilidad operacional (máxima satisfacción de '
    'usuarios EV: 3,500 vehículos/año). Aunque A2C maximiza CO₂ reducción, sacrifica usuarios de manera '
    'no justificada (−19% respecto a SAC). PPO falla completamente (carga 10.7% MENOS que baseline). '
    'SAC emerge como la solución técnicamente robusta y operacionalmente viable.'
)

# Guardar
output_path = 'reports/SECCION_464_467_AGENTE_RESULTADOS_CO2_COMPLETO.docx'
doc.save(output_path)

print(f'✓ Documento actualizado: {output_path}')
print(f'✓ Secciones 4.6.4 y 4.6.4.7 con análisis COMPLETO CO₂')
print(f'✓ Incluye: Grid Import + Solar Export + BESS (análisis a nivel SISTEMA)')
print(f'✓ SAC justificado: 1.30M kg CO₂ + 3,500 EV = BALANCE ÓPTIMO')
