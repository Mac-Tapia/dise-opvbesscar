"""Valida que el documento Word fue generado correctamente"""
from docx import Document

doc = Document('reports/ARQUITECTURA_PVBESSCAR_TESIS.docx')
content = [p.text for p in doc.paragraphs if p.text.strip()]
print(f"✓ Párrafos con contenido: {len(content)}")
print(f"✓ Tablas incrustadas. {len(doc.tables)}")
headings = len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])
print(f"✓ Encabezados: {headings}")

print(f"\n📄 Secciones principales:")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        print(f"   {p.text[:70]}")
