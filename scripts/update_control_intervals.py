#!/usr/bin/env python3
"""
Actualiza documento Word para incluir detalle técnico sobre intervalos de control
y frecuencia de actualización de políticas en cada agente RL.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_oe3_section():
    """Abre documento existente y agrega sección de intervalos de control"""
    doc = Document('reports/ARQUITECTURA_PVBESSCAR_TESIS.docx')
    
    # Buscar la sección "FASE OE3" para insertar contenido después
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if "CRITERIOS DE EVALUACIÓN OE3" in para.text:
            insert_index = i
            break
    
    if insert_index is None:
        print("⚠ No se encontró sección CRITERIOS DE EVALUACIÓN OE3")
        return
    
    # Insertar nueva sección ANTES de CRITERIOS
    # Necesitamos agregar párrafos en el índice correcto
    # Insertaremos una nueva sección de control después de OE3 intro
    
    # Crear documento nuevo limpio e insertar en ubicación correcta
    new_doc = Document()
    
    # Copiar todo hasta el punto de inserción
    for i, para in enumerate(doc.paragraphs):
        if i < insert_index:
            new_para = new_doc.add_paragraph(para.text, style=para.style)
            # Copiar formato
            for run in para.runs:
                if run.text:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
        elif i == insert_index:
            # INSERTAR NUEVA SECCIÓN AQUÍ
            
            new_doc.add_heading('FRECUENCIA DE ACTUALIZACIÓN DE POLÍTICAS POR ALGORITMO', level=3)
            
            intro = new_doc.add_paragraph(
                'Aunque la simulación opera con timesteps horarios (8,760 horas = 1 año simulado), '
                'cada algoritmo RL actualiza internamente su política con frecuencias distintas, '
                'determinadas por su configuración de entrenamiento y naturaleza (on-policy vs off-policy):'
            )
            
            # Tabla de comparación
            tbl = new_doc.add_table(rows=5, cols=5)
            tbl.style = 'Light Grid Accent 1'
            
            # Encabezado
            cells = tbl.rows[0].cells
            cells[0].text = 'Algoritmo'
            cells[1].text = 'Tipo'
            cells[2].text = 'Configuración n_steps'
            cells[3].text = 'Frecuencia Actualización'
            cells[4].text = 'Mecanismo'
            
            # PPO
            cells = tbl.rows[1].cells
            cells[0].text = 'PPO'
            cells[1].text = 'On-Policy'
            cells[2].text = 'n_steps = 2,048'
            cells[3].text = 'Cada 2,048 horas simuladas'
            cells[4].text = 'Acumula 2,048 experiencias en rollout, luego optimiza 10 épocas'
            
            # A2C
            cells = tbl.rows[2].cells
            cells[0].text = 'A2C'
            cells[1].text = 'On-Policy'
            cells[2].text = 'n_steps = 8'
            cells[3].text = 'Cada 8 horas simuladas'
            cells[4].text = 'Actualización frecuente de parámetros cada 8 pasos'
            
            # SAC
            cells = tbl.rows[3].cells
            cells[0].text = 'SAC'
            cells[1].text = 'Off-Policy'
            cells[2].text = 'n_steps = 1'
            cells[3].text = 'Continua (múltiples updates/timestep)'
            cells[4].text = 'Replay buffer (100k) + gradient_steps=1+ por timestep'
            
            # Notas
            cells = tbl.rows[4].cells
            cells[0].text = 'NOTA'
            cells[1].text = ''
            cells[2].text = 'Timestep horario: 1h = 3,600 segundos'
            cells[3].text = 'Cobertura anual: SAC ≥99%, PPO ~85 días, A2C variable'
            cells[4].text = 'Consultar src/agents/{ppo_sb3,a2c_sb3,sac}.py para valores exactos'
            
            new_doc.add_paragraph()
            
            # Detallar cada algoritmo
            new_doc.add_heading('Detalles por Algoritmo', level=4)
            
            # PPO
            new_doc.add_heading('PPO: Actualización cada 2,048 pasos', level=5)
            ppo_text = '''Configuración:
• n_steps = 2,048 pasos (equivalent a ~85 días calendarios de simulación)
• batch_size = 256
• n_epochs = 10 (épocas de optimización por rollout)

Mecánica:
1. Agente acumula 2,048 transiciones (observación → acción → reward)
2. Completa 1 rollout = 2,048 timesteps horarios
3. Optimiza política mediante 10 épocas de descenso de gradiente
4. Vuelve a comenzar rollout nuevo

Implicación: La política se optimiza cada ~85 días simulados. Entre actualizaciones, 
el agente sigue la misma política sin cambios.'''
            
            for line in ppo_text.split('\n'):
                if line.strip():
                    if line.startswith('•'):
                        new_doc.add_paragraph(line, style='List Bullet')
                    elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.'):
                        new_doc.add_paragraph(line, style='List Number')
                    else:
                        new_doc.add_paragraph(line)
            
            new_doc.add_paragraph()
            
            # A2C
            new_doc.add_heading('A2C: Actualización cada 8 pasos', level=5)
            a2c_text = '''Configuración:
• n_steps = 8 pasos (equivalent a 8 horas simuladas / 1 día calendario)
• Learning rate = 7e-4 (relativamente alto para convergencia rápida)

Mecánica:
1. Agente acumula 8 transiciones (observación → acción → reward)
2. Completa 1 mini-rollout = 8 timesteps horarios
3. Actualiza actores/críticos inmediatamente (sin esperar época completa como PPO)
4. Vuelve a comenzar rollout nuevo

Implicación: La política se optimiza cada 8 horas simuladas (~3 veces al día). 
Frecuencia MUCHO mayor que PPO = aprendizaje rápido pero potencialmente inestable.
Fortaleza: A2C es el algoritmo más "reactivo" a cambios en el entorno.'''
            
            for line in a2c_text.split('\n'):
                if line.strip():
                    if line.startswith('•'):
                        new_doc.add_paragraph(line, style='List Bullet')
                    elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.'):
                        new_doc.add_paragraph(line, style='List Number')
                    else:
                        new_doc.add_paragraph(line)
            
            new_doc.add_paragraph()
            
            # SAC
            new_doc.add_heading('SAC: Actualización Continua (Off-Policy)', level=5)
            sac_text = '''Configuración:
• n_steps = 1 (por diseño off-policy; actualiza independientemente de rollout)
• gradient_steps = 1+ (múltiples updates por timestep)
• update_per_time_step = 1+ (permite >1 update/hora)
• buffer_size = 100,000 transiciones (capacidad de 11.4 años de datos)

Mecánica:
1. Agente observa estado, toma acción, recibe reward
2. Almacena transición (s,a,r,s') en replay buffer circular
3. INDEPENDIENTEMENTE de pasos, muestrea batch del buffer
4. Actualiza actor (policy), crítico (Q-values), y entropy coefficient
5. Vuelve a step 1 (sin epocas, sin rollout acumulación)

Implicación: La política se optimiza CONTINUAMENTE en cada timestep, sin esperar 
a que se acumule datos (como PPO) o a minirollouts (como A2C).

Ventaja crítica para este problema:
• SGE simulación anual (8,760 timesteps) → SAC ve panorama anual sin truncamiento
• Buffer de 100k = retiene experiencias antiguas → descubre patrones anuales (estacionalidad solar)
• Actualización continua = adapta rápidamente a cambios (nube, demanda pico)
• Off-policy = puede aprender de trayectorias subóptimas → exploración eficiente'''
            
            for line in sac_text.split('\n'):
                if line.strip():
                    if line.startswith('•'):
                        new_doc.add_paragraph(line, style='List Bullet')
                    elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.') or line.startswith('5.'):
                        new_doc.add_paragraph(line, style='List Number')
                    else:
                        new_doc.add_paragraph(line)
            
            new_doc.add_paragraph()
            
            # Implicaciones en convergencia
            new_doc.add_heading('Implicaciones en Velocidad de Convergencia y Estabilidad', level=4)
            
            convergence_text = '''
RELACIÓN: n_steps ↑ ⟹ convergencia LENTA pero ESTABLE | n_steps ↓ ⟹ convergencia RÁPIDA pero RUIDOSA

| Métrica                    | PPO (2048)  | A2C (8)     | SAC (1)     |
|----------------------------|-------------|-------------|-------------|
| Actualizaciones/año        | 4-5         | ~1,095      | ~87,600+    |
| Datos acumulados/update    | 2,048       | 8           | 1 + buffer  |
| Varianza de gradientes     | BAJA        | MEDIA       | MEDIA-BAJA  |
| Datos anuales capturacos   | ~85 días    | Variable    | 100% anual  |
| Estabilidad               | ALTA        | MEDIA       | MEDIA-ALTA  |
| Capacidad aprendizaje rápido | LENTA     | RÁPIDA      | RÁPIDA      |

CONCLUSIÓN: SAC es óptimo para problemas con horizonte anual porque:
1. Actualización continua → captura variaciones anuales
2. Replay buffer → retiene patrones estacionales
3. Off-policy → explora eficientemente sin esperar rollouts completos

PPO y A2C son mejores para horizontes cortos (energía horaria) pero inferior para panorama anual completo.
'''
            new_doc.add_paragraph(convergence_text)
            
            new_doc.add_paragraph()
            
            # Ahora agregar el resto del documento (CRITERIOS DE EVALUACIÓN en adelante)
            # Agregar párrafo actual (que fue el insert_index)
            new_para = new_doc.add_paragraph(para.text, style=para.style)
            for run in para.runs:
                if run.text:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
        else:
            # Resto del documento después de insertion point
            new_para = new_doc.add_paragraph(para.text, style=para.style)
            for run in para.runs:
                if run.text:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
    
    # Copiar también las tablas del documento original
    for table in doc.tables:
        # Copiar tabla
        new_tbl = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_tbl.style = table.style
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_tbl.rows[i].cells[j].text = cell.text
    
    # Guardar documento actualizado
    new_doc.save('reports/ARQUITECTURA_PVBESSCAR_TESIS.docx')
    print("✓ Documento actualizado con sección de intervalos de control")
    print("✓ Ubicación: reports/ARQUITECTURA_PVBESSCAR_TESIS.docx")

if __name__ == "__main__":
    update_oe3_section()
